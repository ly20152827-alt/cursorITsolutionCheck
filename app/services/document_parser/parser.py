# -*- coding: utf-8 -*-
"""
文档解析模块
支持Word、PDF等格式的文档解析
"""

import os
import re
from typing import Dict, List, Optional
from pathlib import Path


class DocumentParser:
    """文档解析器基类"""
    
    def __init__(self):
        self.supported_formats = []
    
    def parse(self, file_path: str) -> Dict:
        """解析文档"""
        raise NotImplementedError
    
    def extract_chapters(self, content: str) -> List[Dict]:
        """提取章节结构"""
        # 匹配章节标题的正则表达式
        # 支持：1. 标题、第一章 标题、1 标题 等格式
        chapter_patterns = [
            r'^第[一二三四五六七八九十]+章\s+(.+)$',
            r'^第[一二三四五六七八九十]+节\s+(.+)$',
            r'^\d+[\.、]\s*(.+)$',
            r'^\d+\.\d+\s+(.+)$',
            r'^（[一二三四五六七八九十]+）\s*(.+)$',
        ]
        
        chapters = []
        lines = content.split('\n')
        current_chapter = None
        current_section = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # 检查是否是章节标题
            for pattern in chapter_patterns:
                match = re.match(pattern, line)
                if match:
                    title = match.group(1).strip()
                    
                    # 判断是章还是节
                    if '章' in line or re.match(r'^\d+[\.、]', line):
                        current_chapter = {
                            "title": title,
                            "level": 1,
                            "line_number": i + 1,
                            "sections": []
                        }
                        chapters.append(current_chapter)
                        current_section = None
                    elif '节' in line or re.match(r'^\d+\.\d+', line):
                        if current_chapter:
                            current_section = {
                                "title": title,
                                "level": 2,
                                "line_number": i + 1,
                                "content": []
                            }
                            current_chapter["sections"].append(current_section)
                    break
        
        return chapters
    
    def extract_text_content(self, content: str) -> str:
        """提取纯文本内容"""
        # 移除多余空白
        content = re.sub(r'\s+', ' ', content)
        return content.strip()


class WordParser(DocumentParser):
    """Word文档解析器"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.docx', '.doc']
    
    def parse(self, file_path: str) -> Dict:
        """解析Word文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            
            # 提取文本内容
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            content = '\n'.join(full_text)
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # 提取章节结构
            chapters = self.extract_chapters(content)
            
            return {
                "content": content,
                "chapters": chapters,
                "tables": tables,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
        except ImportError:
            raise ImportError("请安装python-docx库: pip install python-docx")
        except Exception as e:
            raise Exception(f"Word文档解析失败: {str(e)}")


class PDFParser(DocumentParser):
    """PDF文档解析器"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.pdf']
    
    def parse(self, file_path: str) -> Dict:
        """解析PDF文档"""
        try:
            import PyPDF2
            
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content.append(page.extract_text())
            
            full_text = '\n'.join(content)
            
            # 提取章节结构
            chapters = self.extract_chapters(full_text)
            
            return {
                "content": full_text,
                "chapters": chapters,
                "page_count": len(pdf_reader.pages)
            }
        except ImportError:
            raise ImportError("请安装PyPDF2库: pip install PyPDF2")
        except Exception as e:
            raise Exception(f"PDF文档解析失败: {str(e)}")


class DocumentParserFactory:
    """文档解析器工厂"""
    
    _parsers = {
        '.docx': WordParser,
        '.doc': WordParser,
        '.pdf': PDFParser,
    }
    
    @classmethod
    def get_parser(cls, file_path: str) -> DocumentParser:
        """根据文件扩展名获取对应的解析器"""
        file_ext = Path(file_path).suffix.lower()
        parser_class = cls._parsers.get(file_ext)
        
        if not parser_class:
            raise ValueError(f"不支持的文件格式: {file_ext}")
        
        return parser_class()
    
    @classmethod
    def parse_document(cls, file_path: str) -> Dict:
        """解析文档"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        parser = cls.get_parser(file_path)
        return parser.parse(file_path)
